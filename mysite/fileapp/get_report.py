import pandas as pd
import numpy as np
import os

import jinja2
from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx.shared import Mm, Inches, Pt
from docx import Document

#生成图表
import matplotlib.pyplot as plt
from docx.shared import Inches
from matplotlib.font_manager import FontProperties


def days_count(df):
    days_this = (pd.to_datetime(df.columns[-2]) - pd.to_datetime(df.columns[-3])).days
    days_all = (pd.to_datetime(df.columns[-2]) - pd.to_datetime(df.columns[-4])).days
    return days_this,days_all

# need update
def wether_alarm(df,alarm_table):
    df['是否预警'] = '否'
    return df

def level_table(context,df,type):
   
    #日期、天数
    days_this,days_all = days_count(df)
    print(days_this,days_all)
    print('***********************************************************')
    df['total'] = (df.iloc[:,3]-df.iloc[:,1])*1000
    df['this'] = (df.iloc[:,3]-df.iloc[:,2])*1000
    df['v'] = df['this']/days_this
    wether_alarm(df,df)
    # print(df)
    
    #将备注列移至最后
    col_to_move = '备注' # 要移动的列名
    last_col = df.pop(col_to_move)  # 使用 pop() 函数将列移除，并返回移除的列
    df.insert(len(df.columns), col_to_move, last_col)  # 使用 insert() 函数将列插入到最后一列位置
    print(df)


    # 将列名更新为和template中一致
    df.columns = ['app_name','app_in','app_last','app_crrt','app_total','app_this','app_v','app_alarm','app_note']

    #求得最大累积和速率
    print(df)
    total_max = df.loc[df['app_total'].abs().idxmax(), 'app_total']
    this_max = df.loc[df['app_this'].abs().idxmax(), 'app_this']
    v_max = df.loc[df['app_v'].abs().idxmax(), 'app_v']
    print(total_max,this_max,v_max)

    #将循环数据转为列名为键，每一行为值的字典
    dict_data = df.to_dict(orient='records')

    #向字典中添加值
    newdata = {'类型':type,
               'total_max':total_max,
               'this_max':this_max,
               'v_max':v_max,
               'app_list':dict_data,
               }
    context.update(newdata)
    # print(context)

    # context = {'user_name': '小明',
    #         'user_sex': '男',
    #         'username': '123456',
    #         'user_birth': '2000年11月',
    #         'user_phone': '123456',
    #         'user_email': '177@qq.com',
    #         'user_birthplace': '山西',
    #         'user_unit': '13班',
    #         'user_photo': InlineImage(tpl, '01.jpg',width=Mm(40),height=Mm(56)),
    #         'app_list': [{'app_name': '奖学金', 'app_data': '2020年11月'}, {'app_name': '荣誉证书', 'app_data': '2020年1月'}]
    #         }

    # tpl.save('template1.docx')
    return context

# 获得表头信息、数据内容df_data(不隐藏的： 备注、测点类型、是否隐藏、点编号、初值……)、测点类型
def pre_level(file):
    df = pd.read_csv(file,delimiter=',') # 读取文件
    df_pretable = df.iloc[0:11,0:3]  #表头信息
    print('###################################3')
    df = df.iloc[:,3:]  #数据信息
    df_data = df[df["是否隐藏"]=="否"]
    df_type = df_data['测点类型'].unique().tolist()
    return df_pretable,df_data,df_type

# 获取用于出报告的信息：点号、初值、上次、本次、备注
def get_data_table_dict(df):
    df = pd.concat([df.iloc[:,3:5], df.iloc[:,-2:], df.iloc[:,0]],axis = 1)
    return df
# 根据数据生成折线图。数据内容df(不隐藏的： 备注、测点类型、是否隐藏、点编号、初值……)
def get_chart_by_type(df,type):
    print(df)
    print('^^^^^^^^^^^^')

    print(df.shape[1])
    ## 曲线图显示90天的曲线
    if df.shape[1] > 95:
        df = pd.concat([df.iloc[:,3:5], df.iloc[:,-90:]],axis=1)
    else:
        df=df.iloc[:,3:]
    print(df)
    # 中文会支持不好
    df.rename(columns={'测点编号': 'Name'}, inplace=True)
    df.set_index('Name',inplace = True)
    #改为零值用前面一次的值填充，对曲线图影响小些
    # 将零值替换为NaN
    df = df.replace(0, np.nan)
    # 前向填充非零值
    df = df.ffill(axis=1) 
    print(df)

    # 创建 Word 文档对象
    doc = Document()

    # # # 设置中文字体
    # font_path = 'SIMSUN.TTC'  # 替换为你的字体文件路径
    # font_prop = FontProperties(fname=font_path)
    # plt.rcParams['font.family'] = font_prop.get_name()
    # 设置绘图风格
    plt. style.use("ggplot")
    # # 设置中文编码和符号的正常显示
    # plt.rcParams["font.sans-serif"] = "KaiTi"
    # plt.rcParams["axes.unicode_minus"] = False

    font = FontProperties(fname='/usr/share/fonts/truetype/ubuntu/Ubuntu-L.ttf', size=15)  # 设置汉字格式
    
    # 将数据按每3行分组，并生成折线图添加到文档中
    groups = [df.iloc[i:i+3] for i in range(0, len(df), 3)]


    for i, group in enumerate(groups):

        # 剔除数值为零的数据, ,
        # group = group.replace(0, pd.NA)
        # group.iloc[:,1:].astype(float)
        print('!!!!!!!!!!!!!!!!!!!!group is')
        print(group) 
        print(group.dtypes)
        if len(group) > 0: 
            # 创建图形对象

            fig, ax = plt.subplots()
            # 绘制折线图
            group.T.plot(ax=ax,
                        linestyle = '-', # 折线类型
                        linewidth = 2, # 折线宽度
                        marker = 'o', # 点的形状
                        markersize = 3, # 点的大小

                         )
            # 添加标题和标签
            plt.title(f'Group {i+1} Line Chart')
            plt.xlabel('Time')
            plt.ylabel('Value')
            # 剔除图框上边界和右边界的刻度
            plt.tick_params( top = 'on', bottom = 'off', right = 'on', left = 'off')
            # 为了避免x轴日期刻度标签的重叠，设置x轴刻度自动展现，并且15度倾斜
            fig.autofmt_xdate(rotation = 15)

            # 设置图例位置为下方居中
            plt.legend(loc='lower center', bbox_to_anchor=(0.5, -0.3), ncol=3)
            # 设置日期的显示格式  
            # date_format = mpl.dates.DateFormatter("%Y-%m-%d")  
            # ax.xaxis.set_major_formatter(date_format) 

            # # 设置x轴显示多少个日期刻度
            # #xlocator = mpl.ticker.LinearLocator(10)
            # # 设置x轴每个刻度的间隔天数
            # xlocator = mpl.ticker.MultipleLocator(5)
            # ax.xaxis.set_major_locator(xlocator)
            # 标注曲线对应的 index
            for index, row in group.iterrows():
                plt.text(len(row)-1, row.iloc[-1], f'Index: {index}')
            # 将图形添加到 Word 文档中
            temp_file = f'line_chart_{i+1}.png'
            print(temp_file)
            plt.savefig(temp_file)
            doc.add_picture(temp_file)
            # 删除临时文件
            plt.close()
            del fig
            del ax

    # 保存 Word 文档
    doc.save('line_charts_'+ type +'.docx')


if __name__=="__main__":
    #上传的文件
    folder_path = 'mysite/fileapp/file_data/'  # 替换为你的文件夹路径
    file = os.path.join(folder_path,'样表.csv')
    df_pretable,df_data,df_type = pre_level(file)

    #处理表头
    print(df_pretable)
    result_dict = {k:v for k,v in zip(df_pretable.iloc[:,1], df_pretable.iloc[:,2])}
    print(result_dict)

    #提取本次数据，处理数据
    # 待添加（前提，确保横轴按时间顺序排列）

    #输入文档表格
    final_doc = Document()
    for type in df_type:
        # 分批次获取用于出报告的数据
        print(type,'*****************************************************************')
        df = get_data_table_dict(df_data[df_data['测点类型'] == type]) #便历不同类型
        # print('$$$$$$$$$')
        # print(df)
        

        #使用模板填充内容
        template = DocxTemplate('mysite/fileapp/file_data/template.docx')
        context = level_table(result_dict,df,type)
        # print(context)
        jinja_env = jinja2.Environment(autoescape=True)
        template.render(context, jinja_env)

        #分类生成结果，合并结果
        template.save(type + '.docx')  # 测试是否成功
      
        type_doc= Document(type + '.docx')
        # 生成折线图，并附在该类表格后面
        get_chart_by_type(df_data[df_data['测点类型'] == type],type)
        doc_line_charts = Document('line_charts_'+ type +'.docx')
        for element in doc_line_charts.element.body:
            print(element)
            type_doc.element.body.append(element)
        type_doc.save(type + '.docx')


        doc = Document(type+'.docx')
        for element in doc.element.body:
            final_doc.element.body.append(element)
        
        # # 在尾部添加
        # filled_doc = template.get_docx()
        # for element in filled_doc.element.body:
        #     final_doc.element.body.append(element)
        #     # final_doc.add_element(element)  
        print('--------------------------------added success-----------------------')
    final_doc.save('final_document.docx')  

